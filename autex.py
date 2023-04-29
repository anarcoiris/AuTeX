import docx
import pandas as pd
import docx2txt

# function to extract all paragraphs with more than 40 characters
def get_text(filename):
    doc = docx.Document(filename)
    full_text = []
    for para in doc.paragraphs:
        if len(para.text) > 40:
            full_text.append("\newline " + para.text)
    return full_text

# function to extract all titles (paragraphs with less than 40 characters)
def get_title(filename):
    doc = docx.Document(filename)
    titles = []
    for para in doc.paragraphs:
        if len(para.text) < 40 and len(para.text) > 3:
            # Add a new page and a chapter tag to each title
            titles.append("\newpage \chapter{" + para.text + "}")
    return titles

# function to divide the document into chapters based on a character limit
def get_chapters(filename, divi):
    # get all the titles first
    titles = get_title(filename)
    chapters = []
    for i, para in enumerate(doc.paragraphs):
        if len(para.text) > 1 and len(para.text) < divi:
            if i == 0 or len(chapters) == 0:
                # add the first chapter with its title
                chapters.append({'title': titles[0], 'text': []})
            else:
                # add the current paragraph to the latest chapter
                chapters[-1]['text'].append("\newline " + para.text)
        elif len(para.text) > divi:
            raise ValueError('Paragraph too long')
    return chapters

# function to dump the chapters into a latex file
def dump_to_latex(filename, chapters):
    with open(filename+'.tex', 'w', encoding='utf-8') as fp:
        for chapter in chapters:
            fp.write(chapter['title'] + '\n')
            fp.write('\n'.join(chapter['text']) + '\n')

if __name__ == '__main__':
    # Prompt the user for the file name
    filename = input("Enter the name of the Word document file (including extension): ")

    # Check if the file exists
    if not os.path.exists(filename):
        print("File not found.")
        sys.exit(1)
    # get the file name and the character limit for dividing the chapters
    divi = 25
    # get the chapters from the file
    chapters = get_chapters(filename, divi)
    # dump the chapters into a latex file
    dump_to_latex(filename, chapters)

    # Read the Word document
    text = docx2txt.process("test.docx")

    # Split the text into sections based on headings
    sections = []
    current_section = {"title": "", "text": ""}
    for line in text.split("\n"):
        if line.startswith("# "):
            if current_section["text"]:
                sections.append(current_section)
                current_section = {"title": "", "text": ""}
            current_section["title"] = line[2:]
        else:
            current_section["text"] += line + "\n"
    sections.append(current_section)

    # Define the template
    template = r"""
    \documentclass{article}
    \begin{document}
    {sections}
    \end{document}
    """

    # Fill in the template
    filled_template = template.format(
        sections="\n\n".join(
            "\\section{{{title}}}\n{text}".format(
                title=section["title"], text=section["text"]
            )
            for section in sections
        )
    )

    # Write the filled template to a file
    with open("document.tex", "w") as f:
        f.write(filled_template)
