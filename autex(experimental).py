#! python3        #PorHacer:  Sensibilidad, arreglar defectos/artefactos. Convertir en fichero de texto con extension .tex
#AutoLateX Templater - s. anarcoiris@pm.me                    ::::===>>> implementar plantillas para completar una primera vrsión!
import docx
import json
   #Prueba para arreglar el formato de las tildes, etc...
   #sys.setdefaultencoding('utf-8')
import sys
#s.encode('utf-8')

Class Transcrib:
   def __init__(self,titulo,texto):
        self.titulo = order["Título"]
        self.texto = order["Texto"]
      
   def readOrden()
        self.text
      
   def read(self,k):
        print(self.titulo[j])
        print(self.texto[k])
      
   keyNum = lastkey + 1
   valueOne=Transcrib(order["Título"],order["Texto"])
   valueTwo
   
   valueOne[j].read()
   valueTwo[k].read()
   
#Class parag:
    #def __init__(self,longt,fuente,color,size,tipo):
        # self.longt=longt
         #self.fuente=fuente
       #  self.color=color
        # self.size=size   
        # self.tipo=tipo   

orden={"Número":[],"Título":[],"Texto":[]}
i = 0
n = 0
Title = []
fullText = []


def getText(filename):
    doc = docx.Document(filename)
    #fullText = []
    for para in doc.paragraphs:
        if len(para.text) > 40:   #Filtro de Title/Body      --- introducir parametro de sensibilidad,
            fullText.append("\newline " + para.text)
    return fullText            #return '\n'.join(fullText)

def getTitle(filename):
    doc = docx.Document(filename)
    #Title = []
    for para in doc.paragraphs:
        if len(para.text) < 40 and len(para.text) > 3:  #Filtro de Title
            orden["Título"].append(" \newpage " + " \chapter{" + para.text + "}")
    return orden #    return Title     return '\n'.join(Title)

def nTitle(filename, n):  #aquí empieza la parte experimental
    doc = docx.Document(filename)
    Title = []
    for para in doc.paragraphs:
        if len(para.text) < 40 and len(para.text) > 3:
            Title.append(" \newpage " + " \chapter{" + para.text + "}")
    print(Title)    #print(Title[(
   )])
    return '\n'.join(Title[n])

def textChain(filename,divi):
    claves(divi)
    getTitle(filename)
    i=0         #Se trata de encadenar los textos que pertenezcan al mismo capítulo.
    doc = docx.Document(filename) #Debería añadir el vector formado por los appends anteriores a la lista cuando encuentre un paragraph de entre menos de 25 y mas de 1
    fullText = []
    nText = []
    Separar=""

    for para in doc.paragraphs:
        if len(para.text) < 4:
            #print("detectado renglón, omitiendo")
            continue
        elif len(para.text) > 40: #   if len(para.text) < 25 and len(para.text) > 1:
            fullText.append(' \newline '+para.text)         #    continue  - Seguimos añadiendo texto a la lista...
        #    print(fullText)
            i+=1
            continue
        elif i>1:                                  #Encontramos un título en nuestro filtro
            #print(fullText)
            x = Separar.join(fullText) #i+=1  ¡¡ La intención es agrupar todos los elementos de la lista en una cadena.
        #    print(fullText)                       Esta contendrá todo el texto entre título y título (EJ Capitulo 1 + párrafos 1,2,3...5) !!
            orden["Texto"].append(x) # Aquí quedan ordenados en una nueva lista, elem 1: párrafos capítulo 1, elem 2: párrafos 2, etc...
            fullText.clear() #Creo que es buena idea limpiar la anterior lista - los párrafos recogidos en la anterior ya no nos sirven - se duplicarían.
            monta=dict(zip(Title,nText))        
            dict(monta)        #Este diccionario está en desuso.... código inútil (probabglemente, tal vez realmente es necesario para el json.dump(dict,...)
            with open('monta.txt', 'w', encoding='utf-8') as fp:
                json.dump(monta, fp, ensure_ascii=False)
            dict(orden)
            with open(filename+'.tex', 'w', encoding='utf-8') as fp:
                json.dump(orden, fp, ensure_ascii=False)
            #print(dict(monta))
            #print([i for i in zip(Title,nText)])
            #n+=1 sin uso - nTitle[n!]
            #print(nText)
            continue
        else:
            i+=1
            continue

        return print(tuple(monta))       # check. cuando termina el for de para.text haría una comprobación de que está montado bien -
    #print(tuple(monta))            el primer texto aparecia en vacío, de ahí la necesidad de la línea 61, asegurando que el título[0] no se quede con text[0]=(void)
    #print([i for i in zip(Title[n],nText)])
    return monta
#    monta = {}
#    f = open("dict.txt","w")             TERRIBLES PRUEBAS
#    f.write( str(dict) )
#    f.close()
#    orden["Título"].append()         ^^ ^^      Esta idea promete. ^^ ^^
    #orden.keys()=Title
    #orden[Title]


def dumpJson(outname,filename,divi):
    textChain(filename,divi)
    monta=dict(zip(Title,nText))
    dict(monta)
    with open('monta.txt', 'w', encoding='utf-8') as fp:
        json.dump(monta, fp, ensure_ascii=False)
    dict(orden)
    with open(filename+'.tex', 'w', encoding='utf-8') as fp:
        json.dump(orden, fp, ensure_ascii=False)

    #print(monta)
            # else:    #if len(para.text) < 25 and len(para.text) > 1:
                    #    break
    #fullText.append("'newpage " + nText)
    #return '\n'.join(fullText)

def getTupla(filename):
    getTitle(filename)
    getText(filename)
    zipped=zip(Title,fullText)
    tuple(zipped)
    print(tuple(zipped))
    #print(list(zip(Title,fullText)))
    print([i for i in zip(Title,fullText)])
    return zipped

def troceo(filename,n):       #Sin uso actualmente
    getTitle(filename)
    getText(filename)
    return print(list(zip(*[iter(Title)]*n)))


def claves(divi):             #Aún no tiene mucha utilidad, pero es algo que he empezado a hacer con previsión de poder acceder facilmente 
    count=0                   #a los títulos y textos correspondientes para poder incluirlos de forma ordenada en las plantillas de LaTeX
    while count < divi:
        orden["Número"].append(count)
        count+=1
    print(orden["Número"])     #check
    return orden["Número"]     
def config(autor,gTitle,instit,edAnar
   
def plantilla(clase)
    #Caso Articulo
    if clase ="Article":
         with open('articletemp.txt', 'w', encoding='utf-8') as fp:
            json.load(orden, fp, ensure_ascii=False)
         return 1
    if clase ="Articulo": 
         return 1
    if clase ="article":    
         return 1
    if clase ="articulo":  
         return 1           
    if clase ="Aritcle":
         return 1
    if clase ="Artcluo":
         return 1
     
    #Caso Libro
    if clase ="Book":
         with open('articletemp.txt', 'w', encoding='utf-8') as fp:
            json.load(orden, fp, ensure_ascii=False)
         return 2
    if clase ="Libro": 
         return 2
    if clase ="book":    
         return 2
    if clase ="libro":  
         return 2           
    if clase ="literatura":
         return 2
    if clase ="ebook":
         return 2
    #Caso Revista
    if clase ="Magazine":
         with open('articletemp.txt', 'w', encoding='utf-8') as fp:
            json.load(orden, fp, ensure_ascii=False)
         return 3
    if clase ="Revista": 
         return 3
    if clase ="magazine":    
         return 3
    if clase ="revista":  
         return 3           
    if clase ="mag":
         return 3
    if clase ="zine":
         return 3      
           
    #Caso Newspaper
    if clase ="News":
         with open('articletemp.txt', 'w', encoding='utf-8') as fp:
            json.load(orden, fp, ensure_ascii=False)
         return 4
    if clase ="Periodico": 
         return 4
    if clase ="news":    
         return 4
    if clase ="periodico":  
         return 4           
    if clase ="newspaper":
         return 4
    if clase ="Newspaper":
         return 4
